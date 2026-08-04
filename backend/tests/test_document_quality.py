"""Task 5: parsing quality gates and source-anchor false-positive boundaries."""

from __future__ import annotations

import importlib
from pathlib import Path

from fastapi.testclient import TestClient

from app.api import documents as documents_api
from app.main import create_app
from app.services.parse_document import parse_document
from app.services.source_chunking import chunk_text


def test_short_chinese_rule_tables_and_symbols_are_accepted(tmp_path: Path):
    path = tmp_path / "short-rule.txt"
    path.write_text(
        "第一条 价格校验规则\n"
        "最小变动单位为 0.001 元；公式 ΔP=(P2-P1)/P1。\n"
        "状态 | 含义\n正常 | SÃO PAULO / ≥ / ≤ / ±",
        encoding="utf-8",
    )

    parsed = parse_document(path)

    assert parsed.diagnostics.quality_ok
    assert not parsed.diagnostics.warnings
    assert "0.001" not in {
        clause_id
        for span in parsed.clause_spans
        for clause_id in span.clause_ids
    }


def test_replacement_character_threshold_warns_before_rejecting(tmp_path: Path):
    warning_path = tmp_path / "warning.txt"
    warning_path.write_text("有效规则" * 100 + "��", encoding="utf-8")
    warning = parse_document(warning_path).diagnostics
    assert warning.quality_ok
    assert warning.replacement_char_count == 2
    assert warning.warnings

    rejected_path = tmp_path / "rejected.txt"
    rejected_path.write_text("有效规则" * 75 + "���", encoding="utf-8")
    rejected = parse_document(rejected_path).diagnostics
    assert not rejected.quality_ok
    assert "替换字符" in " ".join(rejected.errors)


def test_empty_document_is_rejected_before_llm(tmp_app_data, monkeypatch):
    calls = {"count": 0}

    def must_not_call_llm(*args, **kwargs):
        calls["count"] += 1
        raise AssertionError("low-quality input must not reach the LLM")

    monkeypatch.setattr(documents_api, "_INGEST_CHAT_FN", must_not_call_llm)
    client = TestClient(create_app())
    uploaded = client.post(
        "/api/documents",
        files={"file": ("empty.txt", b" \n\t", "text/plain")},
    )
    assert uploaded.status_code == 200

    job = client.post(f"/api/documents/{uploaded.json()['id']}/ingest").json()

    assert job["status"] == "failed"
    assert "未提取到正文" in (job["error_message"] or "")
    assert calls["count"] == 0


class _FakePage:
    def __init__(self, text: str):
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _FakeReader:
    def __init__(self, texts: list[str]):
        self.pages = [_FakePage(text) for text in texts]


def _patch_pdf_reader(monkeypatch, texts: list[str]) -> None:
    parser_module = importlib.import_module("app.services.parse_document")
    monkeypatch.setattr(parser_module, "PdfReader", lambda _path: _FakeReader(texts))


def test_sparse_scanned_pdf_is_rejected_but_substantial_mixed_pdf_is_not(
    tmp_path: Path,
    monkeypatch,
):
    path = tmp_path / "sample.pdf"

    _patch_pdf_reader(monkeypatch, ["封面", "", "", ""])
    sparse = parse_document(path).diagnostics
    assert sparse.suspicious_scanned_pdf
    assert not sparse.quality_ok
    assert "OCR" in " ".join(sparse.errors)

    _patch_pdf_reader(monkeypatch, ["有效正文" * 100, "", ""])
    mixed = parse_document(path).diagnostics
    assert not mixed.suspicious_scanned_pdf
    assert mixed.quality_ok


def test_pdf_page_spans_flow_into_source_chunks(tmp_path: Path, monkeypatch):
    path = tmp_path / "rules.pdf"
    _patch_pdf_reader(
        monkeypatch,
        [
            "第一章 总则\n1.1　适用范围。",
            "第二章 交易\n2.1.1　价格优先、时间优先。",
        ],
    )

    parsed = parse_document(path)
    chunks = chunk_text(parsed, chunk_chars=200, overlap_chars=20)

    assert [span.page_start for span in parsed.page_spans] == [1, 2]
    assert any(chunk["page_start"] == 1 for chunk in chunks)
    assert any(chunk["page_end"] == 2 for chunk in chunks)
    assert {"1.1", "2.1.1"}.issubset(
        {clause_id for chunk in chunks for clause_id in chunk["clause_ids"]}
    )


def test_docx_toc_is_ignored_and_legal_clauses_keep_real_section_anchors(
    tmp_path: Path,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    path = tmp_path / "exchange-rules.docx"
    doc = Document()
    toc_style = doc.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("上海证券交易所交易规则（测试样本）")
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 总则\t- 5 -", style=toc_style)
    doc.add_paragraph("第二章 交易市场\t- 6 -", style=toc_style)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("1.1　为规范证券交易行为，制定本规则。")
    doc.add_paragraph("1.2　价格最小变动单位为0.001元，公式ΔP=(P2-P1)/P1。")
    doc.add_paragraph("第二章 交易市场")
    doc.add_paragraph("第一节 交易场所")
    doc.add_paragraph("2.1.1　证券交易遵循价格优先、时间优先原则。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "申报价格"
    table.cell(1, 1).text = "≥0 且 ≤1000000"
    doc.save(path)

    parsed = parse_document(path)
    chunks = chunk_text(parsed, chunk_chars=300, overlap_chars=30)

    assert parsed.diagnostics.quality_ok
    assert "第一章 总则\t- 5 -" not in parsed.text
    assert [span.section for span in parsed.section_spans] == [
        "第一章 总则",
        "第二章 交易市场",
        "第一节 交易场所",
    ]
    assert {"1.1", "1.2", "2.1.1"}.issubset(
        {clause_id for span in parsed.clause_spans for clause_id in span.clause_ids}
    )
    assert "0.001" not in {
        clause_id for chunk in chunks for clause_id in chunk["clause_ids"]
    }
    assert any(
        chunk["section"] == "第一节 交易场所"
        and "2.1.1" in chunk["clause_ids"]
        for chunk in chunks
    )
    assert not any(chunk["text"].strip() == "第二章 交易市场" for chunk in chunks)
    assert "申报价格" in parsed.text


def test_long_section_does_not_copy_every_clause_id_into_every_chunk(tmp_path: Path):
    path = tmp_path / "long-rules.md"
    path.write_text(
        "# 交易规则\n"
        "1.1　" + ("甲" * 500) + "。\n"
        "1.2　" + ("乙" * 500) + "。\n",
        encoding="utf-8",
    )

    chunks = chunk_text(parse_document(path), chunk_chars=300, overlap_chars=20)
    chunks_for_second_clause = [chunk for chunk in chunks if "乙" * 50 in chunk["text"]]

    assert chunks_for_second_clause
    assert any("1.2" in chunk["clause_ids"] for chunk in chunks_for_second_clause)
    assert all("1.1" not in chunk["clause_ids"] for chunk in chunks_for_second_clause[-1:])
