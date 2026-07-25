from fastapi.testclient import TestClient

from app.main import create_app


def test_upload_txt(tmp_app_data):
    client = TestClient(create_app())
    r = client.post(
        "/api/documents",
        files={"file": ("notes.txt", "余额校验规则".encode("utf-8"), "text/plain")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["status"] == "parsed"
    assert body["char_count"] > 0
    assert body["filename"] == "notes.txt"


def test_upload_empty_markdown_still_parsed(tmp_app_data):
    client = TestClient(create_app())
    r = client.post(
        "/api/documents",
        files={"file": ("empty.md", b"", "text/markdown")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["status"] == "parsed"
    assert body["char_count"] == 0


def test_get_document_not_found(tmp_app_data):
    client = TestClient(create_app())
    assert client.get("/api/documents/123456").status_code == 404


def test_ingest_not_found(tmp_app_data):
    client = TestClient(create_app())
    assert client.post("/api/documents/123456/ingest").status_code == 404


def test_docx_upload_then_ingest_with_fake_llm(tmp_app_data, tmp_path, monkeypatch):
    """Regression: docx must upload+parse and compile end-to-end."""
    from docx import Document

    from app.api import documents as documents_api

    path = tmp_path / "sse.docx"
    doc = Document()
    doc.add_paragraph("上海证券交易所交易规则")
    doc.add_paragraph("集合竞价期间不得撤单")
    doc.save(path)

    def fake_chat(messages):
        system = (messages[0].get("content") or "") if messages else ""
        if "summary_title" in system or "仅 JSON" in system:
            return (
                '{"summary_title":"上交所规则","key_rules":["集合竞价不可撤单"],'
                '"api_points":[],"test_hints":["撤单校验"],'
                '"entities":["集合竞价"],"suggested_page_types":["business"]}'
            )
        return """---
title: 上交所规则摘要
type: source_summary
sources: ["raw/sources/sse.docx"]
tags: ["集合竞价"]
---
集合竞价期间不得撤单。
"""

    monkeypatch.setattr(documents_api, "_INGEST_CHAT_FN", fake_chat)
    client = TestClient(create_app())
    up = client.post(
        "/api/documents",
        files={
            "file": (
                "sse.docx",
                path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code == 200
    assert up.json()["status"] == "parsed"
    assert up.json()["char_count"] > 0
    doc_id = up.json()["id"]

    job = client.post(f"/api/documents/{doc_id}/ingest").json()
    assert job["status"] == "success"
    assert client.get(f"/api/documents/{doc_id}").json()["status"] == "ready"
    pages = client.get("/api/wiki/pages").json()
    assert any("集合竞价" in (p.get("title") or "") or p.get("page_type") for p in pages)
