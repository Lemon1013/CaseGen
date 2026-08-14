from fastapi.testclient import TestClient
from sqlmodel import Session

from app.db import get_engine, init_db
from app.main import create_app
from app.models.entities import Document, IngestJob, SourceChunk, WikiReviewItem
from app.services.wiki_repository import WikiRepository
from app.services.wiki_schema import WikiFrontmatter, WikiPage, WikiSource


def test_upload_markdown(tmp_app_data):
    client = TestClient(create_app())
    content = "# 余额规则\n余额不足应拒绝下单".encode("utf-8")
    files = {"file": ("rules.md", content, "text/markdown")}
    r = client.post("/api/documents", files=files)
    assert r.status_code == 200
    body = r.json()
    assert body["status"] in ("parsed", "uploaded", "ready")
    assert body["filename"] == "rules.md"
    assert body["char_count"] > 0
    assert body["sha256"]
    assert "raw/sources/" in body["stored_path"]

    listed = client.get("/api/documents").json()
    assert len(listed) >= 1
    assert any(item["id"] == body["id"] for item in listed)

    detail = client.get(f"/api/documents/{body['id']}")
    assert detail.status_code == 200
    assert detail.json()["filename"] == "rules.md"


def test_reject_bad_extension(tmp_app_data):
    client = TestClient(create_app())
    files = {"file": ("malware.exe", b"not-a-doc", "application/octet-stream")}
    r = client.post("/api/documents", files=files)
    assert r.status_code == 400
    assert "extension" in r.json()["detail"].lower()


def test_upload_docx(tmp_app_data, tmp_path):
    from docx import Document

    docx_path = tmp_path / "rules.docx"
    doc = Document()
    doc.add_heading("余额规则", level=1)
    doc.add_paragraph("余额不足应拒绝下单")
    doc.save(docx_path)

    client = TestClient(create_app())
    files = {
        "file": (
            "rules.docx",
            docx_path.read_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = client.post("/api/documents", files=files)
    assert r.status_code == 200
    body = r.json()
    assert body["filename"] == "rules.docx"
    assert body["status"] == "parsed"
    assert body["char_count"] > 0


def test_delete_document_archives_or_detaches_wiki_sources(tmp_app_data):
    client = TestClient(create_app())
    first = client.post(
        "/api/documents",
        files={"file": ("first.md", "# 第一份\n规则 A".encode(), "text/markdown")},
    ).json()
    second = client.post(
        "/api/documents",
        files={"file": ("second.md", "# 第二份\n规则 B".encode(), "text/markdown")},
    ).json()
    init_db()
    with Session(get_engine()) as session:
        first_doc = session.get(Document, first["id"])
        assert first_doc is not None and first_doc.space_id is not None
        repository = WikiRepository(session, space_id=first_doc.space_id)
        repository.create(
            WikiPage(
                frontmatter=WikiFrontmatter(
                    page_key="rule.delete.only-source",
                    title="单一来源规则",
                    type="rule",
                    sources=[WikiSource(document_id=first["id"])],
                ),
                body="仅来自第一份文档。",
            )
        )
        repository.create(
            WikiPage(
                frontmatter=WikiFrontmatter(
                    page_key="rule.delete.shared-source",
                    title="共享来源规则",
                    type="rule",
                    sources=[
                        WikiSource(document_id=first["id"]),
                        WikiSource(document_id=second["id"]),
                    ],
                ),
                body="由两份文档共同支持。",
            )
        )
        session.add(
            SourceChunk(
                document_id=first["id"],
                space_id=first_doc.space_id,
                chunk_index=0,
                text="第一份原文块",
            )
        )
        job = IngestJob(
            document_id=first["id"],
            space_id=first_doc.space_id,
            status="success",
        )
        session.add(job)
        session.flush()
        session.add(
            WikiReviewItem(
                job_id=int(job.id),
                space_id=first_doc.space_id,
                kind="review",
                status="pending",
                reason="待核对",
            )
        )
        session.commit()

    source_path = tmp_app_data / first["stored_path"]
    assert source_path.is_file()
    deleted = client.delete(f"/api/documents/{first['id']}")
    assert deleted.status_code == 200, deleted.json()
    body = deleted.json()
    assert body["chunks_deleted"] == 1
    assert body["source_file_removed"] is True
    assert "rule.delete.only-source" in body["pages_archived"]
    assert "rule.delete.shared-source" in body["pages_detached"]
    assert body["reviews_closed"] == 1
    assert not source_path.exists()

    assert all(item["id"] != first["id"] for item in client.get("/api/documents").json())
    assert client.get(f"/api/documents/{first['id']}").status_code == 404
    assert client.get(f"/api/documents/{first['id']}/preview").status_code == 404
    assert client.get(f"/api/documents/{first['id']}/chunks").status_code == 404
    assert client.post(f"/api/documents/{first['id']}/ingest").status_code == 404

    visible_pages = client.get("/api/wiki/pages").json()
    assert {page["page_key"] for page in visible_pages} == {"rule.delete.shared-source"}
    historical_pages = client.get("/api/wiki/pages", params={"include_archived": "true"}).json()
    assert {page["page_key"] for page in historical_pages} == {
        "rule.delete.only-source",
        "rule.delete.shared-source",
    }
    archived_page = next(
        page for page in historical_pages if page["page_key"] == "rule.delete.only-source"
    )
    assert client.get(f"/api/wiki/pages/{archived_page['id']}").status_code == 404
    assert client.get(
        f"/api/wiki/pages/{archived_page['id']}",
        params={"include_archived": "true"},
    ).status_code == 200
    index = client.get("/api/wiki/index").json()["content"]
    assert "rule.delete.only-source" not in index
    assert "共享来源规则" in index
    retrieved = client.post(
        "/api/wiki/retrieve",
        json={"query": "仅来自第一份文档", "top_k": 10},
    )
    assert retrieved.status_code == 200
    assert all(
        hit.get("page_key") != "rule.delete.only-source"
        for hit in retrieved.json()["hits"]
    )

    with Session(get_engine()) as session:
        document = session.get(Document, first["id"])
        assert document is not None and document.status == "deleted"
        repository = WikiRepository(session, space_id=document.space_id)
        assert repository.read("rule.delete.only-source").frontmatter.status == "archived"
        shared = repository.read("rule.delete.shared-source")
        assert [source.document_id for source in shared.frontmatter.sources] == [second["id"]]
        review = session.get(WikiReviewItem, 1)
        assert review is not None and review.status == "rejected"


def test_delete_document_blocks_active_ingest(tmp_app_data):
    client = TestClient(create_app())
    document = client.post(
        "/api/documents",
        files={"file": ("busy.md", b"# busy", "text/markdown")},
    ).json()
    with Session(get_engine()) as session:
        row = session.get(Document, document["id"])
        assert row is not None
        session.add(
            IngestJob(
                document_id=document["id"],
                space_id=row.space_id,
                status="queued",
            )
        )
        session.commit()

    response = client.delete(f"/api/documents/{document['id']}")
    assert response.status_code == 409
    assert client.get(f"/api/documents/{document['id']}").status_code == 200
