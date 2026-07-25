from pathlib import Path

import pytest

from app.services.parse_document import parse_file


def test_parse_markdown_and_txt(tmp_path: Path):
    md = tmp_path / "a.md"
    md.write_text("# 标题\n余额不足", encoding="utf-8")
    assert "余额不足" in parse_file(md)

    txt = tmp_path / "b.txt"
    txt.write_text("纯文本规则", encoding="utf-8")
    assert parse_file(txt) == "纯文本规则"


def test_parse_docx_paragraphs_and_tables(tmp_path: Path):
    from docx import Document

    path = tmp_path / "rules.docx"
    doc = Document()
    doc.add_heading("交易规则", level=1)
    doc.add_paragraph("集合竞价在 9:15 开始")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "证券代码"
    table.cell(1, 1).text = "6 位"
    doc.save(path)

    text = parse_file(path)
    assert "交易规则" in text
    assert "集合竞价" in text
    assert "证券代码" in text
    assert "6 位" in text


def test_parse_unsupported_extension(tmp_path: Path):
    bad = tmp_path / "x.exe"
    bad.write_bytes(b"MZ")
    with pytest.raises(ValueError, match="Unsupported extension"):
        parse_file(bad)
