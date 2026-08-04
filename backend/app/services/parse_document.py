"""Document parsing with lossless character and source-location metadata.

``parse_file`` remains a small compatibility wrapper returning only text.  New
callers should use ``parse_document`` so page/section anchors and extraction
quality are retained for SourceChunk and Wiki citations.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from pypdf import PdfReader


@dataclass(frozen=True)
class TextSpan:
    """A half-open character range with optional source-location metadata."""

    start_char: int
    end_char: int
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    section: Optional[str] = None
    clause_ids: tuple[str, ...] = ()

    @property
    def start(self) -> int:
        return self.start_char

    @property
    def end(self) -> int:
        return self.end_char


@dataclass
class QualityDiagnostics:
    """Deterministic, user-readable extraction quality signals."""

    is_empty: bool = False
    replacement_char_count: int = 0
    replacement_char_rate: float = 0.0
    garbled_char_count: int = 0
    garbled_char_rate: float = 0.0
    suspicious_scanned_pdf: bool = False
    page_count: int = 0
    pages_with_text: int = 0
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def quality_ok(self) -> bool:
        return not self.errors

    @property
    def ok(self) -> bool:
        return self.quality_ok

    @property
    def messages(self) -> list[str]:
        return [*self.errors, *self.warnings]


@dataclass
class ParsedDocument:
    """Parsed text plus ranges needed to cite the original source."""

    text: str
    page_spans: list[TextSpan] = field(default_factory=list)
    section_spans: list[TextSpan] = field(default_factory=list)
    clause_spans: list[TextSpan] = field(default_factory=list)
    diagnostics: QualityDiagnostics = field(default_factory=QualityDiagnostics)

    @property
    def quality_diagnostics(self) -> QualityDiagnostics:
        """Descriptive alias used by callers that prefer an explicit name."""

        return self.diagnostics

    @property
    def spans(self) -> list[TextSpan]:
        """Combined spans, useful for simple integrations."""

        return [*self.page_spans, *self.section_spans, *self.clause_spans]


_MD_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*#*\s*$")
_CN_SECTION_HEADING_RE = re.compile(
    r"^\s*(?:第[0-9一二三四五六七八九十百千万]+[编章节目]|附则)"
    r"(?:[\s\u3000:：、.-]+.*)?\s*$"
)
_NUMERIC_HEADING_RE = re.compile(
    r"^\s*[1-9]\d*(?:\.\d+){0,2}[\s\u3000、.．:：-]+(.+?)\s*$"
)
_CLAUSE_LINE_RE = re.compile(
    r"(?m)^[ \t]*(?P<id>"
    r"[1-9]\d*(?:\.\d+){1,3}|"
    r"第[0-9一二三四五六七八九十百千万]+条"
    r")(?=[ \t\u3000、.．:：-])"
)
_MOJIBAKE_SEQUENCES = ("Ã¤", "Ã¥", "Ã©", "â€", "â€™", "â€œ", "ðŸ")


def _unique_clause_ids(text: str) -> tuple[str, ...]:
    seen: set[str] = set()
    result: list[str] = []
    for match in _CLAUSE_LINE_RE.finditer(text or ""):
        value = match.group("id")
        if value not in seen:
            seen.add(value)
            result.append(value)
    return tuple(result)


def _clause_spans(
    text: str,
    section_spans: list[TextSpan] | None = None,
) -> list[TextSpan]:
    """Return article ranges without mistaking prices or formulas for clauses."""

    matches = list(_CLAUSE_LINE_RE.finditer(text or ""))
    section_starts = sorted(span.start_char for span in (section_spans or []))
    spans: list[TextSpan] = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        next_section = next((value for value in section_starts if value > start), None)
        if next_section is not None:
            end = min(end, next_section)
        spans.append(
            TextSpan(
                start_char=start,
                end_char=end,
                clause_ids=(match.group("id"),),
            )
        )
    return spans


def _heading_label(line: str) -> str | None:
    """Recognise structural headings, not ordinary numbered legal clauses."""

    markdown = _MD_HEADING_RE.match(line)
    if markdown:
        return markdown.group(1).strip()
    stripped = line.strip()
    if _CN_SECTION_HEADING_RE.match(stripped):
        return stripped
    numeric = _NUMERIC_HEADING_RE.match(stripped)
    if not numeric:
        return None
    # A short numbered label such as ``2.1 API`` is a useful section.  A legal
    # sentence such as ``2.1 本规则适用于……。`` remains a clause anchor only.
    tail = numeric.group(1).strip()
    if len(stripped) <= 80 and not re.search(r"[。；：:！？!?]$", tail):
        return stripped
    return None


def _diagnose(
    text: str,
    *,
    suffix: str,
    page_count: int = 0,
    pages_with_text: int = 0,
) -> QualityDiagnostics:
    non_space = max(1, sum(not ch.isspace() for ch in text))
    replacement_count = text.count("\ufffd")
    mojibake_count = sum(text.count(value) for value in _MOJIBAKE_SEQUENCES)
    control_count = sum(
        unicodedata.category(ch) == "Cc" and ch not in "\n\r\t\f"
        for ch in text
    )
    garbled_count = replacement_count + mojibake_count + control_count
    replacement_rate = replacement_count / non_space
    garbled_rate = garbled_count / non_space
    suspicious_pdf = False
    if suffix == ".pdf":
        # A blank/textless PDF is the reliable scan signal.  For multi-page
        # PDFs, a majority of textless pages is also a useful conservative
        # signal.  Do not reject a legitimate short one-page text PDF.
        sparse_text = sum(not ch.isspace() for ch in text) < max(200, page_count * 40)
        suspicious_pdf = page_count > 0 and (
            pages_with_text == 0
            or (page_count >= 3 and pages_with_text * 2 < page_count and sparse_text)
        )

    diagnostics = QualityDiagnostics(
        is_empty=not text.strip(),
        replacement_char_count=replacement_count,
        replacement_char_rate=round(replacement_rate, 6),
        garbled_char_count=garbled_count,
        garbled_char_rate=round(garbled_rate, 6),
        suspicious_scanned_pdf=suspicious_pdf,
        page_count=page_count,
        pages_with_text=pages_with_text,
    )
    if diagnostics.is_empty:
        diagnostics.errors.append("未提取到正文内容，请检查文件是否为空或需要 OCR")
    # A small amount of damage is reported but does not kill a reasonable
    # short note.  High replacement/garbling density is not safe as evidence.
    if replacement_rate >= 0.02 or (replacement_count >= 3 and replacement_rate >= 0.005):
        diagnostics.errors.append(
            f"文本包含过多替换字符（{replacement_count} 个，比例 {replacement_rate:.1%}）"
        )
    elif replacement_count:
        diagnostics.warnings.append(
            f"文本包含 {replacement_count} 个替换字符，部分内容可能无法还原"
        )
    if garbled_rate >= 0.12 and garbled_count >= 3:
        diagnostics.errors.append(
            f"疑似乱码（异常字符比例 {garbled_rate:.1%}），请检查文件编码"
        )
    if suspicious_pdf:
        diagnostics.errors.append("PDF 未提取到足够文本，疑似扫描件，请先进行 OCR")
    return diagnostics


def _section_spans(text: str) -> list[TextSpan]:
    """Identify Markdown and common Chinese/numeric section headings."""

    headings: list[tuple[int, str]] = []
    offset = 0
    for line in text.splitlines(keepends=True):
        raw = line.rstrip("\r\n")
        label = _heading_label(raw)
        if label and len(label) <= 200:
            headings.append((offset, label))
        offset += len(line)
    spans: list[TextSpan] = []
    for index, (start, label) in enumerate(headings):
        end = headings[index + 1][0] if index + 1 < len(headings) else len(text)
        if end > start:
            spans.append(
                TextSpan(
                    start_char=start,
                    end_char=end,
                    section=label,
                    clause_ids=_unique_clause_ids(text[start:end]),
                )
            )
    return spans


def _parse_plain(path: Path) -> ParsedDocument:
    raw = path.read_bytes()
    text = raw.decode("utf-8-sig", errors="replace")
    sections = _section_spans(text)
    return ParsedDocument(
        text=text,
        section_spans=sections,
        clause_spans=_clause_spans(text, sections),
        diagnostics=_diagnose(text, suffix=path.suffix.lower()),
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    reader = PdfReader(str(path))
    page_parts: list[str] = []
    page_spans: list[TextSpan] = []
    cursor = 0
    pages_with_text = 0
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_with_text += 1
        if page_parts:
            page_parts.append("\n")
            cursor += 1
        start = cursor
        page_parts.append(page_text)
        cursor += len(page_text)
        page_spans.append(
            TextSpan(
                start_char=start,
                end_char=cursor,
                page_start=page_number,
                page_end=page_number,
                clause_ids=_unique_clause_ids(page_text),
            )
        )
    text = "".join(page_parts).strip()
    # Keep offsets valid after stripping only leading/trailing whitespace.
    leading = len("".join(page_parts)) - len("".join(page_parts).lstrip())
    trailing = len("".join(page_parts).rstrip())
    if leading or trailing != len("".join(page_parts)):
        page_spans = [
            TextSpan(
                start_char=max(0, span.start_char - leading),
                end_char=max(0, min(len(text), span.end_char - leading)),
                page_start=span.page_start,
                page_end=span.page_end,
                clause_ids=span.clause_ids,
            )
            for span in page_spans
            if span.end_char - leading > 0 and span.start_char - leading < len(text)
        ]
    diagnostics = _diagnose(
        text,
        suffix=".pdf",
        page_count=len(reader.pages),
        pages_with_text=pages_with_text,
    )
    sections = _section_spans(text)
    return ParsedDocument(
        text=text,
        page_spans=page_spans,
        section_spans=sections,
        clause_spans=_clause_spans(text, sections),
        diagnostics=diagnostics,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document  # python-docx

    doc = Document(str(path))
    parts: list[str] = []
    heading_offsets: list[tuple[int, str]] = []
    cursor = 0

    def add_line(value: str, *, heading: bool = False) -> None:
        nonlocal cursor
        value = value.strip()
        if not value:
            return
        if parts:
            parts.append("\n")
            cursor += 1
        start = cursor
        parts.append(value)
        cursor += len(value)
        if heading:
            heading_offsets.append((start, value))

    for para in doc.paragraphs:
        value = para.text or ""
        style = (getattr(para.style, "name", "") or "").lower()
        # Word TOC paragraphs duplicate headings and poison source anchors.
        if style.startswith("toc") or style.startswith("目录"):
            continue
        is_heading = "heading" in style or bool(_CN_SECTION_HEADING_RE.match(value.strip()))
        add_line(value, heading=is_heading)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            add_line(" | ".join(c for c in cells if c))
    text = "".join(parts).strip()
    sections: list[TextSpan] = []
    # Word heading styles and explicit Chinese chapter/section labels are
    # reliable. Numeric body paragraphs remain clause anchors, not sections.
    if heading_offsets:
        sections = []
        for i, (start, label) in enumerate(heading_offsets):
            end = heading_offsets[i + 1][0] if i + 1 < len(heading_offsets) else len(text)
            sections.append(
                TextSpan(start, end, section=label, clause_ids=_unique_clause_ids(text[start:end]))
            )
    return ParsedDocument(
        text=text,
        section_spans=sections,
        clause_spans=_clause_spans(text, sections),
        diagnostics=_diagnose(text, suffix=".docx"),
    )


def parse_document(path: Path | str) -> ParsedDocument:
    """Parse a supported document and retain source anchors and diagnostics."""

    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _parse_plain(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported extension: {suffix}")


def parse_file(path: Path | str) -> str:
    """Backward-compatible text-only parser."""

    return parse_document(path).text
